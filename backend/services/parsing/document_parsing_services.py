import fitz  # PyMuPDF
import docx
import os
import logging
# Lazy import for easyocr - only load when needed
# import easyocr
from PIL import Image
import io
import json
from datetime import datetime
import hashlib
import base64

# Import Gemini service for AI-powered text extraction
try:
    from backend.services.ai_services.gemini_service import GeminiService
    GEMINI_AVAILABLE = True
except ImportError:
    logging.warning("Gemini service not available, will use OCR fallback")
    GEMINI_AVAILABLE = False
    GeminiService = None

# Configure logging with UTF-8 encoding
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('document_parser.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)

class DocumentParser:
    def __init__(self):
        """Initialize the document parser with AI and OCR support"""
        # Lazy load OCR reader - only initialize when needed
        self.ocr_reader = None
        self._ocr_initialized = False

        # Initialize Gemini service for AI-powered text extraction
        self.gemini_service = GeminiService() if GEMINI_AVAILABLE else None
        if self.gemini_service and not self.gemini_service.is_available():
            logging.warning("Gemini service initialized but not available")
            self.gemini_service = None

        # Supported file extensions
        self.supported_extensions = {
            '.pdf': self.parse_pdf,
            '.docx': self.parse_docx,
            '.doc': self.parse_docx,  # Will try to handle .doc files
            '.jpg': self.parse_image,
            '.jpeg': self.parse_image,
            '.png': self.parse_image,
            '.bmp': self.parse_image,
            '.tiff': self.parse_image,
            '.webp': self.parse_image,
            '.txt': self.parse_text_file
        }
    
    def _init_ocr(self):
        """Lazy initialization of EasyOCR reader"""
        if self._ocr_initialized:
            return
        
        try:
            import easyocr
            self.ocr_reader = easyocr.Reader(['en'], gpu=False)
            logging.info("✅ EasyOCR initialized successfully")
        except Exception as e:
            logging.error(f"❌ Failed to initialize EasyOCR: {e}")
            self.ocr_reader = None
        
        self._ocr_initialized = True
    
    def validate_file(self, file_path):
        """Validate if file exists and is supported"""
        if not os.path.exists(file_path):
            return False, f"File not found: {file_path}"
        
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.supported_extensions:
            return False, f"Unsupported file type: {ext}"
        
        # Check file size (max 25MB)
        file_size = os.path.getsize(file_path)
        max_size = 25 * 1024 * 1024  # 25MB
        if file_size > max_size:
            return False, f"File too large: {file_size / (1024*1024):.1f}MB (max 25MB)"
        
        return True, "File is valid"
    
    def get_file_hash(self, file_path):
        """Get MD5 hash of file for caching"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def parse_text_file(self, file_path):
        """Parse plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return {
                'text': content,
                'word_count': len(content.split()),
                'char_count': len(content),
                'pages': 1,
                'extraction_method': 'text_file'
            }
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
                return {
                    'text': content,
                    'word_count': len(content.split()),
                    'char_count': len(content),
                    'pages': 1,
                    'extraction_method': 'text_file_latin1'
                }
            except Exception as e:
                return {'error': f'Failed to read text file: {str(e)}'}
        except Exception as e:
            return {'error': f'Error parsing text file: {str(e)}'}
    
    def parse_docx(self, file_path):
        """Parse DOCX files with enhanced error handling"""
        try:
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            text_content = '\n'.join(paragraphs)
            
            # Extract tables if any
            tables_content = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append(' | '.join(row_text))
                tables_content.append('\n'.join(table_text))
            
            if tables_content:
                text_content += '\n\nTABLES:\n' + '\n\n'.join(tables_content)
            
            result = {
                'text': text_content,
                'word_count': len(text_content.split()),
                'char_count': len(text_content),
                'paragraphs': len(paragraphs),
                'tables': len(tables_content),
                'extraction_method': 'python_docx'
            }
            
            logging.info(f"✅ DOCX parsed: {len(paragraphs)} paragraphs, {len(tables_content)} tables")
            return result
            
        except Exception as e:
            logging.error(f"❌ Error parsing DOCX: {e}")
            return {'error': f'Failed to parse DOCX: {str(e)}'}
    
    def parse_pdf(self, file_path, use_ai_extraction=True):
        """Parse PDF with AI-powered text extraction for images - Priority: BLIP → LLaVA → OCR"""
        try:
            text_content = ""
            ai_content = ""
            total_pages = 0
            images_processed = 0

            doc = fitz.open(file_path)
            total_pages = len(doc)

            logging.info(f"📄 Starting PDF parsing: {os.path.basename(file_path)} ({total_pages} pages)")

            for page_num, page in enumerate(doc):
                # Extract regular text
                page_text = page.get_text()
                if page_text.strip():
                    text_content += f"\n--- Page {page_num + 1} Text ---\n"
                    text_content += page_text

                # Extract images and apply AI extraction if enabled
                if use_ai_extraction:
                    images = page.get_images(full=True)
                    for img_index, img in enumerate(images):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]

                            # Priority: BLIP (fast local) → LLaVA (local AI) → OCR (last resort)
                            extracted_text = None

                            # 1. Try BLIP first (fast local model)
                            import tempfile
                            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                                temp_file.write(image_bytes)
                                temp_path = temp_file.name

                            try:
                                from services.ai_services.vision_service import VisionService
                                vision_service = VisionService(use_low_res=True, max_image_size=512)
                                logging.info(f"🎨 Extracting from PDF page {page_num + 1}, image {img_index + 1} using BLIP (priority)")
                                blip_result = vision_service.process_image(temp_path, tasks=["caption", "scene"])
                                
                                blip_text = ""
                                if blip_result.get("caption"):
                                    blip_text += blip_result["caption"] + " "
                                if blip_result.get("scene"):
                                    scene_info = blip_result["scene"]
                                    if isinstance(scene_info, dict):
                                        blip_text += " ".join([str(v) for v in scene_info.values() if v])
                                
                                if blip_text and len(blip_text.strip()) > 0:
                                    ai_content += f"\n--- Page {page_num + 1}, Image {img_index + 1} BLIP Extracted ---\n"
                                    ai_content += blip_text.strip()
                                    images_processed += 1
                                    logging.info(f"✅ BLIP extracted text from PDF image {img_index + 1}")
                                    os.unlink(temp_path)
                                    continue  # Success, skip to next image
                            except Exception as blip_error:
                                logging.warning(f"⚠️ BLIP failed for PDF image {img_index + 1}: {blip_error}")

                            # 2. Try LLaVA vision model (second fallback)
                            try:
                                from services.ai_services.ai_agents.ollama_service import ollama_service
                                logging.info(f"🖼️ Extracting from PDF page {page_num + 1}, image {img_index + 1} using LLaVA (fallback)")
                                extracted_text = ollama_service.process_image(temp_path)
                                if extracted_text and len(extracted_text.strip()) > 0:
                                    ai_content += f"\n--- Page {page_num + 1}, Image {img_index + 1} LLaVA Extracted ---\n"
                                    ai_content += extracted_text
                                    images_processed += 1
                                    logging.info(f"✅ LLaVA extracted text from PDF image {img_index + 1}")
                                    os.unlink(temp_path)
                                    continue  # Success, skip to next image
                            except Exception as llava_error:
                                logging.warning(f"⚠️ LLaVA failed for PDF image {img_index + 1}: {llava_error}")
                            finally:
                                if os.path.exists(temp_path):
                                    os.unlink(temp_path)

                            # 3. OCR as absolute last resort (only if both AI methods failed)
                            logging.warning(f"⚠️ Both BLIP and LLaVA failed for PDF image, using OCR as last resort")
                            self._init_ocr()
                            if self.ocr_reader:
                                try:
                                    # Convert bytes to PIL Image for OCR
                                    image = Image.open(io.BytesIO(image_bytes))
                                    ocr_result = self.ocr_reader.readtext(image, detail=0)

                                    if ocr_result:
                                        ocr_text = " ".join(ocr_result)
                                        ai_content += f"\n--- Page {page_num + 1}, Image {img_index + 1} OCR Last Resort ---\n"
                                        ai_content += ocr_text
                                        images_processed += 1
                                        logging.info(f"✅ OCR last resort extracted text from PDF image {img_index + 1}")
                                except Exception as ocr_error:
                                    logging.error(f"❌ OCR last resort failed for PDF image {img_index + 1}: {ocr_error}")

                        except Exception as img_error:
                            logging.warning(f"⚠️ Error processing image {img_index} on page {page_num}: {img_error}")
                            continue

            doc.close()

            # Combine text and AI/OCR content
            combined_text = text_content
            if ai_content:
                combined_text += "\n\n=== AI EXTRACTED TEXT ===\n" + ai_content

            result = {
                'text': combined_text,
                'text_only': text_content,
                'ai_content': ai_content,
                'word_count': len(combined_text.split()),
                'char_count': len(combined_text),
                'pages': total_pages,
                'images_processed': images_processed,
                'extraction_method': 'pymupdf_with_ai' if use_ai_extraction else 'pymupdf_text_only'
            }

            logging.info(f"✅ PDF parsed: {total_pages} pages, {images_processed} images processed with AI/OCR")
            return result

        except Exception as e:
            logging.error(f"❌ Error parsing PDF {file_path}: {e}")
            return {'error': f'Failed to parse PDF: {str(e)}'}
    
    def parse_image(self, file_path):
        """Parse images with priority: BLIP (fast local) → LLaVA (local AI) → OCR (last resort)"""
        try:
            image = Image.open(file_path)
            result = {'image_size': image.size, 'image_mode': image.mode}

            # 1. Try BLIP first (fast local model for image captioning and text)
            try:
                from services.ai_services.vision_service import VisionService
                vision_service = VisionService(use_low_res=True, max_image_size=512)
                logging.info(f"🎨 Attempting BLIP for: {os.path.basename(file_path)} (priority)")
                blip_result = vision_service.process_image(file_path, tasks=["caption", "scene"])
                
                # Extract text from BLIP caption and scene analysis
                blip_text = ""
                if blip_result.get("caption"):
                    blip_text += blip_result["caption"] + " "
                if blip_result.get("scene"):
                    scene_info = blip_result["scene"]
                    if isinstance(scene_info, dict):
                        blip_text += " ".join([str(v) for v in scene_info.values() if v])
                
                if blip_text and len(blip_text.strip()) > 20:
                    result.update({
                        'text': blip_text.strip(),
                        'word_count': len(blip_text.split()),
                        'char_count': len(blip_text),
                        'extraction_method': 'blip_vision',
                        'metadata': blip_result
                    })
                    logging.info(f"✅ BLIP extracted {len(blip_text)} characters")
                    return result
                logging.warning("⚠️ BLIP returned insufficient content")
            except Exception as e:
                logging.warning(f"⚠️ BLIP failed: {str(e)}")

            # 2. Try LLaVA vision model (second fallback - local AI with OCR capability)
            try:
                from services.ai_services.ai_agents.ollama_service import ollama_service
                logging.info(f"🖼️ Attempting LLaVA for: {os.path.basename(file_path)} (fallback)")
                llava_text = ollama_service.process_image(file_path)
                if llava_text and len(llava_text.strip()) > 20:
                    result.update({
                        'text': llava_text,
                        'word_count': len(llava_text.split()),
                        'char_count': len(llava_text),
                        'extraction_method': 'llava_vision'
                    })
                    logging.info(f"✅ LLaVA extracted {len(llava_text)} characters")
                    return result
            except Exception as e:
                logging.warning(f"⚠️ LLaVA failed: {str(e)}")

            # 3. OCR as absolute last resort (only if both AI methods failed)
            logging.warning(f"⚠️ Both BLIP and LLaVA failed for image, using OCR as last resort")
            self._init_ocr()
            if self.ocr_reader:
                try:
                    ocr_result = self.ocr_reader.readtext(file_path, detail=0)
                    if ocr_result:
                        text_content = " ".join(ocr_result)
                        result.update({
                            'text': text_content,
                            'word_count': len(text_content.split()),
                            'char_count': len(text_content),
                            'extraction_method': 'ocr_last_resort'
                        })
                        logging.info(f"✅ OCR last resort extracted {len(text_content)} characters")
                        return result
                except Exception as e:
                    logging.error(f"❌ OCR last resort failed: {str(e)}")

            # If all methods fail, return empty result with metadata
            result.update({
                'text': '',
                'word_count': 0,
                'char_count': 0,
                'extraction_method': 'no_text_extracted',
                'error': 'All extraction methods failed'
            })
            return result

        except Exception as e:
            logging.error(f"❌ Image parsing failed: {str(e)}")
            return {
                'error': f'Failed to parse image: {str(e)}',
                'extraction_method': 'failed'
            }

    def parse_document(self, file_path, use_multimodal_pdf=True, cache_results=True):
        """Main parsing function with caching and validation"""

        # Validate file
        is_valid, message = self.validate_file(file_path)
        if not is_valid:
            return {'error': message}
        ext = os.path.splitext(file_path)[1].lower()
        file_info = {
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'file_size': os.path.getsize(file_path),
            'file_extension': ext,
            'processed_at': datetime.utcnow().isoformat(),
            'file_hash': self.get_file_hash(file_path)
        }

        try:
            # Priority: BLIP (fast local) → LLaVA (local AI) → OCR (last resort) for images
            if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.webp']:
                # 1. Try BLIP first (fast local model)
                try:
                    from services.ai_services.vision_service import VisionService
                    vision_service = VisionService(use_low_res=True, max_image_size=512)
                    logging.info(f"🎨 Using BLIP for image: {os.path.basename(file_path)} (priority)")
                    blip_result = vision_service.process_image(file_path, tasks=["caption", "scene"])
                    
                    blip_text = ""
                    if blip_result.get("caption"):
                        blip_text += blip_result["caption"] + " "
                    if blip_result.get("scene"):
                        scene_info = blip_result["scene"]
                        if isinstance(scene_info, dict):
                            blip_text += " ".join([str(v) for v in scene_info.values() if v])
                    
                    if blip_text and len(blip_text.strip()) > 20:
                        result = {
                            'text': blip_text.strip(),
                            'word_count': len(blip_text.split()),
                            'char_count': len(blip_text),
                            'file_path': file_path,
                            'file_name': os.path.basename(file_path),
                            'file_size': os.path.getsize(file_path),
                            'file_extension': ext,
                            'processed_at': datetime.utcnow().isoformat(),
                            'extraction_method': 'blip_vision',
                            'metadata': blip_result
                        }
                        result.update(file_info)
                        logging.info(f"✅ BLIP extracted {len(blip_text)} characters from image")
                        return result
                except Exception as e:
                    logging.warning(f"⚠️ BLIP failed for image: {e}")

                # 2. Try LLaVA if BLIP failed
                try:
                    from services.ai_services.ai_agents.ollama_service import ollama_service
                    logging.info(f"🖼️ Using LLaVA for image: {os.path.basename(file_path)} (fallback)")
                    llava_text = ollama_service.process_image(file_path)
                    if llava_text and len(llava_text.strip()) > 0:
                        result = {
                            'text': llava_text,
                            'word_count': len(llava_text.split()),
                            'char_count': len(llava_text),
                            'file_path': file_path,
                            'file_name': os.path.basename(file_path),
                            'file_size': os.path.getsize(file_path),
                            'file_extension': ext,
                            'processed_at': datetime.utcnow().isoformat(),
                            'extraction_method': 'llava_vision'
                        }
                        result.update(file_info)
                        logging.info(f"✅ LLaVA extracted {len(llava_text)} characters from image")
                        return result
                except Exception as e:
                    logging.warning(f"⚠️ LLaVA failed for image: {e}")

                # 3. OCR as last resort
                logging.warning("⚠️ Both BLIP and LLaVA failed for image, using OCR as last resort")
                return self._parse_image_fallback(file_path)
            else:
                # For other file types, use existing parse_document logic
                parser_func = self.supported_extensions.get(ext)
                if parser_func:
                    result = parser_func(file_path)
                    result.update(file_info)
                    if cache_results and 'error' not in result:
                        self.cache_result(file_path, result)
                    return result
                else:
                    return {'error': f'Unsupported file type: {ext}'}

        except Exception as e:
            logging.error(f"❌ Unexpected error parsing {file_path}: {e}")
            return {
                'error': f'Unexpected error: {str(e)}',
                **file_info
            }

    def _parse_image_fallback(self, file_path):
        """Fallback image parsing using OCR or other methods"""
        try:
            self._init_ocr()
            if not self.ocr_reader:
                return {'error': 'OCR reader not available for fallback image parsing.'}
            ocr_result = self.ocr_reader.readtext(file_path, detail=0)
            if ocr_result:
                text_content = " ".join(ocr_result)
                image = Image.open(file_path)
                result = {
                    'text': text_content,
                    'word_count': len(text_content.split()),
                    'char_count': len(text_content),
                    'image_size': image.size,
                    'image_mode': image.mode,
                    'extraction_method': 'easyocr_fallback'
                }
                logging.info(f"✅ OCR fallback completed: {len(text_content)} characters extracted")
                return result
            else:
                return {
                    'text': '',
                    'word_count': 0,
                    'char_count': 0,
                    'message': 'No text found in image',
                    'extraction_method': 'no_text_found'
                }
        except Exception as e:
            logging.error(f"❌ Error in fallback image parsing: {e}")
            return {'error': f'Failed fallback image parsing: {str(e)}'}

    def _use_gemini_ocr(self, file_path, is_handwritten=False):
        """Direct Gemini OCR method for testing purposes"""
        try:
            if not self.gemini_service:
                return None

            # Create prompt based on handwriting
            if is_handwritten:
                prompt = "Extract all handwritten text from this image. Transcribe it accurately, preserving the original formatting as much as possible. Return only the text content without any additional commentary."
            else:
                prompt = "Extract all visible text from this image. Return only the text content without any additional commentary."

            logging.info(f"🤖 Using Gemini OCR for: {os.path.basename(file_path)}")
            extracted_text = self.gemini_service.extract_text_from_image(file_path, prompt)

            if extracted_text and len(extracted_text.strip()) > 0:
                logging.info(f"✅ Gemini OCR extracted {len(extracted_text)} characters")
                return extracted_text.strip()
            else:
                logging.warning("⚠️ Gemini OCR returned empty result")
                return None

        except Exception as e:
            logging.error(f"❌ Error in Gemini OCR: {e}")
            return None

    def cache_result(self, file_path, result):
        """Cache parsing results for faster repeated access"""
        try:
            cache_dir = os.path.join(os.path.dirname(__file__), '..', '..', 'cache', 'parsing')
            os.makedirs(cache_dir, exist_ok=True)
            
            cache_file = os.path.join(cache_dir, f"{result['file_hash']}.json")
            
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
            
            logging.info(f"✅ Cached parsing result: {cache_file}")
        except Exception as e:
            logging.warning(f"⚠️ Failed to cache result: {e}")
    
    def get_cached_result(self, file_path):
        """Get cached parsing result if available"""
        try:
            file_hash = self.get_file_hash(file_path)
            cache_file = os.path.join(os.path.dirname(__file__), '..', '..', 'cache', 'parsing', f"{file_hash}.json")
            
            if os.path.exists(cache_file):
                with open(cache_file, 'r', encoding='utf-8') as f:
                    cached_result = json.load(f)
                
                logging.info(f"✅ Using cached result for: {os.path.basename(file_path)}")
                return cached_result
            
            return None
        except Exception as e:
            logging.warning(f"⚠️ Failed to get cached result: {e}")
            return None

# Master parser function (compatible with existing code)
def master_parser(file_path, use_multimodal_pdf=True, use_cache=True):
    """
    Master function to parse different file types.

    Args:
        file_path (str): Path to the file to parse
        use_multimodal_pdf (bool): Whether to use AI extraction for PDF images
        use_cache (bool): Whether to use cached results

    Returns:
        dict: Parsed content with metadata
    """
    parser = DocumentParser()
    # Check cache first
    if use_cache:
        cached_result = parser.get_cached_result(file_path)
        if cached_result:
            logging.info(f"📋 Using cached result for {os.path.basename(file_path)}")
            return cached_result

    # Parse document
    logging.info(f"🔍 Starting parsing for {os.path.basename(file_path)}")
    result = parser.parse_document(file_path, use_multimodal_pdf, use_cache)

    if 'error' in result:
        logging.error(f"❌ Parsing failed for {os.path.basename(file_path)}: {result['error']}")
    else:
        logging.info(f"✅ Parsing completed for {os.path.basename(file_path)}: {result.get('word_count', 0)} words")

    return result

# Testing function
def test_parser_comprehensive():
    """Comprehensive testing function for the parser"""
    parser = DocumentParser()
    
    test_files = [
        # Add test files here
        # 'test_files/sample.pdf',
        # 'test_files/sample.docx',
        # 'test_files/sample.jpg'
    ]
    
    results = []
    for file_path in test_files:
        if os.path.exists(file_path):
            print(f"\n🧪 Testing: {file_path}")
            result = parser.parse_document(file_path)
            results.append(result)
            
            if 'error' in result:
                print(f"❌ Error: {result['error']}")
            else:
                print(f"✅ Success: {result.get('word_count', 0)} words extracted")
                print(f"📄 Method: {result.get('extraction_method', 'unknown')}")
        else:
            print(f"⚠️ Test file not found: {file_path}")
    
    return results

# Utility functions
def get_supported_formats():
    """Get list of supported file formats"""
    parser = DocumentParser()
    return list(parser.supported_extensions.keys())

def estimate_processing_time(file_path):
    """Estimate processing time based on file size and type"""
    if not os.path.exists(file_path):
        return "File not found"
    
    file_size = os.path.getsize(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    
    # Base time estimates (in seconds)
    base_times = {
        '.txt': 0.1,
        '.docx': 0.5,
        '.pdf': 2.0,  # Without OCR
        '.jpg': 1.0, '.jpeg': 1.0, '.png': 1.0,
        '.bmp': 1.5, '.tiff': 2.0, '.webp': 1.0
    }
    
    base_time = base_times.get(ext, 1.0)
    
    # Adjust for file size (rough estimate)
    size_mb = file_size / (1024 * 1024)
    estimated_time = base_time * (1 + size_mb * 0.5)
    
    # Add OCR time for PDFs
    if ext == '.pdf':
        estimated_time += size_mb * 2  # Additional OCR time
    
    return f"{estimated_time:.1f} seconds"

if __name__ == "__main__":
    # Example usage and testing
    print("🔧 Document Parser Service")
    print(f"📋 Supported formats: {', '.join(get_supported_formats())}")
    
    # Run tests if test files exist
    test_parser_comprehensive()